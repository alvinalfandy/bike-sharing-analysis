"""Generate Laporan Proyek Akhir Praktikum Data Mining 2026 (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size, bold=True)
    return p


def add_subheading(doc, text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size, bold=True)
    return p


def add_body(doc, text, justify=True, size=12):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r, size)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size)
    return p


def add_number(doc, text, size=12):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r, size)
    return p


doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ============ HALAMAN JUDUL ============
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN PROYEK AKHIR PRAKTIKUM DATA MINING 2026")
set_font(r, 14, bold=True)

for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analisis dan Prediksi Penyewaan Sepeda Menggunakan Beberapa Algoritma Machine Learning Berbasis Streamlit")
set_font(r, 13, bold=True)

doc.add_paragraph()
# Placeholder Logo Universitas
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[ Sisipkan Logo Universitas Pelita Bangsa di sini ]")
set_font(r, 11, italic=True)
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Disusun oleh :")
set_font(r, 12)

anggota = [
    "1. Alvin Alfandy (312310473)",
    "2. Abidzar Sabil Handoyo (312310471)",
    "3. Ridho Fauzi (312310463)",
]
for a in anggota:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(a)
    set_font(r, 12)

for _ in range(6):
    doc.add_paragraph()
for line in ["PROGRAM STUDI TEKNIK INFORMATIKA", "FAKULTAS TEKNIK", "UNIVERSITAS PELITA BANGSA", "BEKASI", "2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    set_font(r, 12, bold=True)

doc.add_page_break()

# ============ LEMBAR PENGESAHAN ============
add_heading(doc, "LEMBAR PENGESAHAN PROYEK")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN PROYEK AKHIR PRAKTIKUM DATA MINING 2026")
set_font(r, 12, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analisis dan Prediksi Penyewaan Sepeda Menggunakan Beberapa Algoritma Machine Learning Berbasis Streamlit")
set_font(r, 12, bold=True, italic=True)
doc.add_paragraph()
add_body(doc, "Laporan proyek akhir ini telah disusun dan diselesaikan oleh:")
for a in anggota:
    p = doc.add_paragraph()
    r = p.add_run(a)
    set_font(r, 12)
doc.add_paragraph()
add_body(doc, "Sebagai salah satu syarat untuk menyelesaikan Praktikum Data Mining pada Program Studi Teknik Informatika, Fakultas Teknik, Universitas Pelita Bangsa.")
doc.add_paragraph()
add_body(doc, "Laporan ini telah diperiksa dan disetujui untuk disahkan oleh Dosen Pengampu.", justify=False)
p = doc.add_paragraph()
r = p.add_run("Dosen Pengampu")
set_font(r, 12)
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Agung Nugroho, S.Kom., M.Kom.")
set_font(r, 12, bold=True)

doc.add_page_break()

# ============ KATA PENGANTAR ============
add_heading(doc, "KATA PENGANTAR")
doc.add_paragraph()
add_body(doc, "Segala puji dan syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas limpahan rahmat dan karunia-Nya, sehingga kami dapat menyelesaikan laporan proyek akhir praktikum yang berjudul \u201cAnalisis dan Prediksi Penyewaan Sepeda Menggunakan Beberapa Algoritma Machine Learning Berbasis Streamlit\u201d ini dengan baik dan lancar.")
add_body(doc, "Laporan ini disusun sebagai bentuk dokumentasi dan pemahaman kami atas materi yang telah dipelajari dalam Praktikum Data Mining, khususnya dalam mengimplementasikan algoritma data mining ke dalam aplikasi berbasis web interaktif menggunakan Streamlit. Melalui proyek ini, kami mendapatkan pengalaman langsung dalam mengolah dataset, menerapkan algoritma regresi dan klasifikasi, serta memvisualisasikan hasilnya secara dinamis.")
add_body(doc, "Dalam penyusunan laporan ini, kami menyadari bahwa keberhasilan yang dicapai tidak lepas dari bantuan, dukungan, dan arahan dari berbagai pihak. Oleh karena itu, dengan penuh rasa hormat dan terima kasih, kami menyampaikan apresiasi kepada:")
add_number(doc, "Bapak Agung Nugroho, S.Kom., M.Kom., selaku dosen pengampu praktikum, atas bimbingan, arahan, dan ilmu yang telah diberikan selama proses pembelajaran.")
add_number(doc, "Rekan-rekan mahasiswa di kelas Praktikum Data Mining 2026 yang turut berbagi pengetahuan dan pengalaman selama sesi praktikum.")
add_number(doc, "Semua pihak yang telah membantu, baik secara langsung maupun tidak langsung, dalam penyusunan laporan dan pengembangan proyek ini.")
add_body(doc, "Kami menyadari bahwa laporan ini masih memiliki keterbatasan. Oleh karena itu, kami sangat terbuka terhadap kritik dan saran yang membangun demi penyempurnaan di masa yang akan datang. Semoga laporan ini dapat memberikan manfaat bagi pembaca serta menjadi referensi tambahan dalam penerapan data mining secara praktis dan aplikatif.")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Bekasi, 10 Juli 2026")
set_font(r, 12)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Kelompok 1")
set_font(r, 12)

doc.add_page_break()

# ============ DAFTAR ISI ============
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

add_heading(doc, "DAFTAR ISI")
doc.add_paragraph()
daftar_isi = [
    ("LEMBAR PENGESAHAN PROYEK", "i", True, False),
    ("KATA PENGANTAR", "ii", True, False),
    ("DAFTAR ISI", "iii", True, False),
    ("BAB I PENDAHULUAN", "1", True, False),
    ("1.1 Latar Belakang", "1", False, True),
    ("1.2 Tujuan", "1", False, True),
    ("1.3 Manfaat", "2", False, True),
    ("BAB II TINJAUAN PUSTAKA", "3", True, False),
    ("2.1 Tinjauan Jurnal Pertama", "3", False, True),
    ("2.2 Tinjauan Jurnal Kedua", "3", False, True),
    ("2.3 Tinjauan Jurnal Ketiga", "4", False, True),
    ("BAB III IMPLEMENTASI", "5", True, False),
    ("3.1 Arsitektur Aplikasi", "5", False, True),
    ("3.2 Dataset", "5", False, True),
    ("3.3 Alur Sistem", "6", False, True),
    ("3.4 Kode Program", "6", False, True),
    ("BAB IV HASIL DAN PEMBAHASAN", "7", True, False),
    ("4.1 Tampilan Aplikasi", "7", False, True),
    ("4.2 Evaluasi Model", "8", False, True),
    ("4.3 Visualisasi Prediksi", "9", False, True),
    ("4.4 Feature Importance", "9", False, True),
    ("BAB V PENUTUP", "11", True, False),
    ("5.1 Kesimpulan", "11", False, True),
    ("5.2 Saran", "11", False, True),
    ("DAFTAR PUSTAKA", "12", True, False),
    ("LAMPIRAN", "13", True, False),
]
for text, page, bold, indent in daftar_isi:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.left_indent = Inches(0.3)
    pf.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text + "\t" + page)
    set_font(r, 12, bold=bold)

doc.add_page_break()

# ============ BAB I ============
add_heading(doc, "BAB I")
add_heading(doc, "PENDAHULUAN")
doc.add_paragraph()
add_subheading(doc, "1.1 Latar Belakang")
add_body(doc, "Sepeda merupakan salah satu moda transportasi yang semakin populer di berbagai kota besar, terutama dengan hadirnya sistem penyewaan sepeda (bike sharing). Sistem ini memungkinkan masyarakat menyewa sepeda secara fleksibel untuk perjalanan singkat. Namun, pengelolaan armada sepeda menjadi tantangan tersendiri karena jumlah permintaan penyewaan sangat dipengaruhi oleh berbagai faktor seperti cuaca, musim, hari kerja, serta waktu dalam sehari.")
add_body(doc, "Ketidakseimbangan antara jumlah sepeda yang tersedia dengan permintaan pengguna dapat menyebabkan penurunan kualitas layanan. Oleh karena itu, dibutuhkan sebuah pendekatan prediktif yang mampu memperkirakan jumlah penyewaan sepeda berdasarkan kondisi tertentu, sehingga penyedia layanan dapat melakukan optimasi distribusi armada.")
add_body(doc, "Kemajuan teknologi ilmu data dan machine learning memberikan peluang besar dalam pengembangan aplikasi prediksi yang cepat dan akurat. Dengan memanfaatkan berbagai algoritma seperti Linear Regression, Random Forest, Support Vector Machine, K-Nearest Neighbors, dan Decision Tree, pola penyewaan sepeda dapat dimodelkan secara efektif. Melalui framework Streamlit, model tersebut dapat disajikan dalam bentuk aplikasi web interaktif yang memberikan hasil prediksi secara real-time beserta visualisasi data yang mudah dipahami.")
add_subheading(doc, "1.2 Tujuan")
add_number(doc, "Mengembangkan aplikasi berbasis web untuk memprediksi jumlah penyewaan sepeda menggunakan beberapa algoritma machine learning.")
add_number(doc, "Menerapkan dan membandingkan algoritma Linear Regression, Random Forest, SVM, KNN, dan Decision Tree dalam menganalisis dataset penyewaan sepeda.")
add_number(doc, "Menyediakan visualisasi dan hasil prediksi secara interaktif menggunakan framework Streamlit.")
add_subheading(doc, "1.3 Manfaat")
add_number(doc, "Memberikan sarana prediksi jumlah penyewaan sepeda yang mudah diakses dan digunakan oleh penyedia layanan maupun pihak terkait.")
add_number(doc, "Membantu optimasi pengelolaan dan distribusi armada sepeda berdasarkan pendekatan berbasis data.")
add_number(doc, "Menjadi contoh penerapan praktis algoritma data mining dalam menyelesaikan permasalahan nyata di bidang transportasi.")

doc.add_page_break()

# ============ BAB II ============
add_heading(doc, "BAB II")
add_heading(doc, "TINJAUAN PUSTAKA")
doc.add_paragraph()
add_subheading(doc, "2.1 Tinjauan Jurnal Pertama")
add_body(doc, "Penelitian yang dilakukan oleh Yu-Chun Yin, Chi-Shuen Lee, dan Yu-Po Wong dari Stanford University dalam jurnal berjudul \u201cDemand Prediction of Bicycle Sharing Systems\u201d bertujuan untuk memprediksi jumlah penyewaan sepeda per jam berdasarkan kondisi pada jam tersebut. Penelitian ini menggunakan dataset Capital Bikeshare yang dikumpulkan di area Washington D.C. pada tahun 2011 dan 2012, yang juga tersedia pada UCI Machine Learning Repository dengan 17.379 baris data. Empat algoritma dibandingkan, yaitu Ridge Linear Regression, Support Vector Regression (SVR), Random Forest, dan Gradient Boosted Tree.")
add_body(doc, "Hasil penelitian menunjukkan bahwa nilai Root Mean Squared Logarithmic Error (RMSLE) pada data uji mencapai 0,82 untuk Ridge Linear Regression, 0,34 untuk SVR, serta 0,31 untuk Random Forest dan Gradient Boosted Tree setelah dilakukan optimasi parameter dan seleksi fitur. Penelitian ini menegaskan bahwa metode berbasis pohon seperti Random Forest memberikan performa terbaik dan menjadi dasar kuat dalam pengembangan sistem prediksi penyewaan sepeda menggunakan dataset yang sama dengan proyek ini [1].")
add_subheading(doc, "2.2 Tinjauan Jurnal Kedua")
add_body(doc, "Penelitian oleh Jaume Torres, Enrique Jim\u00e9nez-Mero\u00f1o, dan Francesc Soriguera dari Universitat Polit\u00e8cnica de Catalunya dalam jurnal \u201cForecasting the Usage of Bike-Sharing Systems through Machine Learning Techniques to Foster Sustainable Urban Mobility\u201d membahas prediksi penggunaan sistem bike sharing untuk mendukung mobilitas perkotaan yang berkelanjutan. Penelitian ini dilatarbelakangi oleh masalah distribusi sepeda yang tidak seimbang, sehingga operator perlu melakukan repositioning yang mahal. Studi kasus dilakukan pada sistem bike sharing New York City.")
add_body(doc, "Tiga teknik regresi machine learning dibandingkan, yaitu Random Forest, Gradient Boosting, dan Artificial Neural Network, dengan input berupa data historis penggunaan dan data meteorologi. Hasil penelitian menunjukkan bahwa ketiga metode memiliki akurasi yang serupa, namun proses kalibrasi Random Forest yang lebih sederhana membuatnya paling direkomendasikan untuk sebagian besar aplikasi. Temuan ini memperkuat pemilihan Random Forest sebagai algoritma utama dalam proyek ini [2].")
add_subheading(doc, "2.3 Tinjauan Jurnal Ketiga")
add_body(doc, "Penelitian oleh Chang Gao dan Yong Chen dalam jurnal \u201cUsing Machine Learning Methods to Predict Demand for Bike Sharing\u201d menerapkan empat model machine learning, yaitu Linear Regression, K-Nearest Neighbors (KNN), Random Forest, dan Support Vector Machine, untuk memprediksi permintaan bike sharing di Seoul. Penelitian ini memanfaatkan 29 fitur dalam enam kategori, termasuk cuaca, polusi udara, informasi lalu lintas, dan kasus Covid-19.")
add_body(doc, "Hasil penelitian menunjukkan bahwa dua model terbaik adalah Random Forest dan Support Vector Machine. Fitur pada kategori cuaca, polusi, dan wabah Covid-19 merupakan variabel paling penting dalam prediksi model. Penelitian ini menegaskan bahwa pemilihan algoritma dan fitur yang tepat sangat berpengaruh terhadap akurasi prediksi permintaan sepeda, serta menjadi referensi dalam penggunaan beberapa algoritma sekaligus seperti pada proyek ini [3].")

doc.add_page_break()

# ============ BAB III ============
add_heading(doc, "BAB III")
add_heading(doc, "IMPLEMENTASI")
doc.add_paragraph()
add_subheading(doc, "3.1 Arsitektur Aplikasi")
add_body(doc, "Aplikasi ini dirancang untuk menganalisis dan memprediksi jumlah penyewaan sepeda berdasarkan faktor cuaca, musim, dan waktu, menggunakan bahasa pemrograman Python dan framework front-end Streamlit. Model prediksi yang digunakan berasal dari pustaka scikit-learn, meliputi Linear Regression, Random Forest, Support Vector Machine (SVM/SVR), K-Nearest Neighbors (KNN), dan Decision Tree. Visualisasi data didukung oleh pustaka Plotly, Matplotlib, dan Seaborn. Aplikasi ini juga menyertakan fitur interaktif seperti upload dataset, slider input untuk prediksi, evaluasi model, serta visualisasi feature importance.")
add_subheading(doc, "3.2 Dataset")
add_body(doc, "Dataset yang digunakan adalah Bike Sharing Dataset yang diperoleh dari UCI Machine Learning Repository. Dataset ini terdiri dari 17.379 baris data dan 17 kolom yang merekam data penyewaan sepeda per jam pada tahun 2011 hingga 2012. Atribut yang digunakan sebagai fitur meliputi:")
fitur_ds = [
    "season: musim (1: semi, 2: panas, 3: gugur, 4: dingin)",
    "yr: tahun (0: 2011, 1: 2012)",
    "mnth: bulan (1 sampai 12)",
    "hr: jam (0 sampai 23)",
    "holiday: hari libur (0: bukan, 1: ya)",
    "weekday: hari dalam seminggu (0 sampai 6)",
    "workingday: hari kerja (0: bukan, 1: ya)",
    "weathersit: situasi cuaca (1: cerah sampai 4: ekstrem)",
    "temp: suhu ternormalisasi",
    "atemp: suhu terasa ternormalisasi",
    "hum: kelembaban ternormalisasi",
    "windspeed: kecepatan angin ternormalisasi",
]
for f in fitur_ds:
    add_bullet(doc, f)
add_body(doc, "Variabel target adalah cnt (jumlah total penyewaan sepeda). Untuk analisis klasifikasi, nilai cnt dikelompokkan menjadi dua kategori, yaitu Tinggi dan Rendah, berdasarkan nilai median.")
add_subheading(doc, "3.3 Alur Sistem")
add_number(doc, "Pengguna mengunggah file dataset (.csv) atau menggunakan dataset default.")
add_number(doc, "Data ditampilkan dalam bentuk tabel dan statistik ringkas.")
add_number(doc, "Sistem melakukan eksplorasi data melalui visualisasi tren, korelasi, dan distribusi.")
add_number(doc, "Fitur dan target dipisahkan (fitur: kolom numerik; target: cnt).")
add_number(doc, "Dataset dibagi menjadi 80% data latih dan 20% data uji.")
add_number(doc, "Model machine learning dilatih menggunakan algoritma yang dipilih pengguna.")
add_number(doc, "Model diuji dengan data uji dan dievaluasi menggunakan R2, RMSE, MAE (regresi) atau Accuracy, F1, Precision, Recall (klasifikasi).")
add_number(doc, "Hasil prediksi ditampilkan dalam bentuk angka dan visualisasi (scatter plot, confusion matrix, feature importance).")
add_number(doc, "Pengguna dapat melakukan prediksi interaktif dengan memasukkan nilai input melalui slider.")
add_subheading(doc, "3.4 Kode Program")
add_body(doc, "Kode program lengkap dari aplikasi ini dapat diakses melalui repository GitHub pada tautan berikut:")
p = doc.add_paragraph()
r = p.add_run("https://github.com/alvinalfandy/bike-sharing-analysis")
set_font(r, 12)
r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_page_break()

# ============ BAB IV ============
add_heading(doc, "BAB IV")
add_heading(doc, "HASIL DAN PEMBAHASAN")
doc.add_paragraph()
add_subheading(doc, "4.1 Tampilan Aplikasi")
add_body(doc, "Aplikasi menampilkan antarmuka interaktif yang terdiri dari beberapa bagian utama, yaitu Eksplorasi Data, Pemodelan, Prediksi, dan Data Mentah. Pada bagian Eksplorasi Data, pengguna dapat melihat tren penyewaan sepeda per jam, matriks korelasi antar fitur, serta distribusi masing-masing fitur numerik. Pada bagian Pemodelan, pengguna dapat memilih tipe analisis (regresi atau klasifikasi), algoritma, serta fitur yang digunakan untuk melatih model.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[ Sisipkan Gambar 1. Tampilan Dashboard Aplikasi ]")
set_font(r, 11, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gambar 1. Tampilan Dashboard Analisis Bike Sharing")
set_font(r, 11)
add_body(doc, "Gambar 1 menampilkan tampilan awal aplikasi yang berisi statistik ringkas dataset seperti total baris, total kolom, rata-rata penyewaan, dan total penyewaan. Bagian ini memberikan gambaran umum kepada pengguna mengenai data yang sedang dianalisis.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[ Sisipkan Gambar 2. Tampilan Prediksi Interaktif ]")
set_font(r, 11, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gambar 2. Tampilan Fitur Prediksi Interaktif")
set_font(r, 11)
add_body(doc, "Gambar 2 menunjukkan fitur prediksi interaktif, di mana pengguna dapat memasukkan nilai berbagai variabel melalui slider seperti suhu, kelembaban, kecepatan angin, musim, jam, dan hari kerja. Setelah menekan tombol Hitung Prediksi, aplikasi akan menampilkan estimasi jumlah penyewaan sepeda beserta kategori statusnya.")

add_subheading(doc, "4.2 Evaluasi Model")
add_body(doc, "Model dievaluasi menggunakan algoritma Random Forest yang memberikan performa terbaik. Hasil evaluasi untuk analisis regresi dan klasifikasi ditunjukkan pada tabel berikut:")

# Tabel evaluasi regresi
p = doc.add_paragraph()
r = p.add_run("Tabel 1. Hasil Evaluasi Model Regresi (Random Forest)")
set_font(r, 11, bold=True)
t1 = doc.add_table(rows=5, cols=2)
t1.style = "Table Grid"
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
data1 = [("Metrik", "Nilai"), ("R2 Score (Train)", "0,9925 (99,25%)"), ("R2 Score (Test)", "0,9448 (94,48%)"), ("RMSE", "41,80"), ("MAE", "24,72")]
for i, (a, b) in enumerate(data1):
    c0, c1 = t1.rows[i].cells
    r0 = c0.paragraphs[0].add_run(a); set_font(r0, 11, bold=(i == 0))
    r1 = c1.paragraphs[0].add_run(b); set_font(r1, 11, bold=(i == 0))

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Tabel 2. Hasil Evaluasi Model Klasifikasi (Random Forest)")
set_font(r, 11, bold=True)
t2 = doc.add_table(rows=6, cols=2)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
data2 = [("Metrik", "Nilai"), ("Accuracy (Train)", "1,0000 (100%)"), ("Accuracy (Test)", "0,9252 (92,52%)"), ("F1 Score", "0,9245"), ("Precision", "0,9165"), ("Recall", "0,9326")]
for i, (a, b) in enumerate(data2):
    c0, c1 = t2.rows[i].cells
    r0 = c0.paragraphs[0].add_run(a); set_font(r0, 11, bold=(i == 0))
    r1 = c1.paragraphs[0].add_run(b); set_font(r1, 11, bold=(i == 0))

doc.add_paragraph()
add_body(doc, "Interpretasi:")
add_bullet(doc, "Nilai R2 Test sebesar 0,9448 menunjukkan bahwa model mampu menjelaskan 94,48% variasi jumlah penyewaan sepeda berdasarkan fitur input, yang berarti model memiliki performa sangat baik.")
add_bullet(doc, "Nilai akurasi test sebesar 92,52% pada klasifikasi menandakan model mampu mengelompokkan tingkat permintaan (tinggi/rendah) dengan sangat akurat.")
add_bullet(doc, "Nilai RMSE dan MAE yang relatif kecil dibandingkan skala data menandakan error prediksi yang rendah.")

add_subheading(doc, "4.3 Visualisasi Prediksi")
add_body(doc, "Scatter plot memperlihatkan sebaran nilai aktual (y_test) terhadap nilai prediksi (y_pred), dengan garis acuan y = x. Titik-titik yang mendekati garis tersebut mengindikasikan prediksi yang akurat. Selain itu, aplikasi juga menampilkan residual plot untuk menganalisis pola error, serta confusion matrix untuk mengevaluasi hasil klasifikasi.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[ Sisipkan Gambar 3. Scatter Plot Aktual vs Prediksi ]")
set_font(r, 11, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gambar 3. Visualisasi Aktual vs Prediksi")
set_font(r, 11)

add_subheading(doc, "4.4 Feature Importance")
add_body(doc, "Analisis feature importance dari model Random Forest menunjukkan tingkat kontribusi masing-masing fitur terhadap hasil prediksi. Hasil analisis ditunjukkan pada tabel berikut:")
p = doc.add_paragraph()
r = p.add_run("Tabel 3. Feature Importance Model Random Forest")
set_font(r, 11, bold=True)
imp_data = [("Fitur", "Importance"), ("hr (jam)", "0,6118"), ("temp (suhu)", "0,1189"), ("yr (tahun)", "0,0809"), ("workingday (hari kerja)", "0,0584"), ("hum (kelembaban)", "0,0265"), ("atemp (suhu terasa)", "0,0217"), ("season (musim)", "0,0215"), ("weathersit (cuaca)", "0,0181"), ("mnth (bulan)", "0,0165"), ("weekday (hari)", "0,0126"), ("windspeed (angin)", "0,0104"), ("holiday (libur)", "0,0026")]
t3 = doc.add_table(rows=len(imp_data), cols=2)
t3.style = "Table Grid"
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b) in enumerate(imp_data):
    c0, c1 = t3.rows[i].cells
    r0 = c0.paragraphs[0].add_run(a); set_font(r0, 11, bold=(i == 0))
    r1 = c1.paragraphs[0].add_run(b); set_font(r1, 11, bold=(i == 0))
doc.add_paragraph()
add_body(doc, "Berdasarkan Tabel 3, fitur hr (jam) merupakan faktor paling dominan dengan nilai importance sebesar 0,6118. Hal ini menunjukkan bahwa waktu dalam sehari sangat memengaruhi jumlah penyewaan sepeda, di mana jam-jam sibuk seperti pagi dan sore hari cenderung memiliki permintaan tinggi. Faktor temp (suhu) menempati posisi kedua dengan importance 0,1189, yang berarti kondisi suhu turut berperan penting dalam menentukan minat pengguna untuk menyewa sepeda.")

doc.add_page_break()

# ============ BAB V ============
add_heading(doc, "BAB V")
add_heading(doc, "PENUTUP")
doc.add_paragraph()
add_subheading(doc, "5.1 Kesimpulan")
add_body(doc, "Berdasarkan hasil perancangan dan implementasi aplikasi prediksi penyewaan sepeda menggunakan beberapa algoritma machine learning, maka dapat disimpulkan beberapa hal sebagai berikut:")
add_number(doc, "Aplikasi yang dikembangkan berhasil melakukan pemrosesan dataset, pelatihan model prediktif, serta menampilkan evaluasi dan hasil prediksi dengan antarmuka yang interaktif menggunakan Streamlit.")
add_number(doc, "Algoritma Random Forest menunjukkan performa prediksi terbaik dengan nilai R2 Test sebesar 94,48% untuk regresi dan akurasi 92,52% untuk klasifikasi, sehingga memenuhi kriteria model dengan performa sangat baik.")
add_number(doc, "Faktor waktu (jam) dan suhu merupakan variabel yang paling berpengaruh terhadap jumlah penyewaan sepeda berdasarkan analisis feature importance.")
add_number(doc, "Penggunaan Streamlit sebagai framework antarmuka sangat membantu dalam menyederhanakan proses interaksi pengguna, sehingga pengguna non-programmer pun dapat melakukan prediksi dengan mudah.")
add_subheading(doc, "5.2 Saran")
add_number(doc, "Disarankan untuk menambahkan variabel eksternal lain yang relevan, seperti data event khusus atau kondisi lalu lintas, guna meningkatkan akurasi prediksi.")
add_number(doc, "Performa model dapat ditingkatkan lebih lanjut dengan menerapkan teknik hyperparameter tuning dan validasi silang (cross-validation).")
add_number(doc, "Pengembangan lebih lanjut dapat mencakup fitur peramalan deret waktu (time series forecasting) untuk memprediksi permintaan sepeda pada periode mendatang, serta integrasi dengan sistem manajemen armada secara real-time.")

doc.add_page_break()

# ============ DAFTAR PUSTAKA ============
add_heading(doc, "DAFTAR PUSTAKA")
doc.add_paragraph()
pustaka = [
    "[1] Y. C. Yin, C. S. Lee, and Y. P. Wong, \u201cDemand Prediction of Bicycle Sharing Systems,\u201d CS229 Project Report, Stanford University, 2014.",
    "[2] J. Torres, E. Jim\u00e9nez-Mero\u00f1o, and F. Soriguera, \u201cForecasting the Usage of Bike-Sharing Systems through Machine Learning Techniques to Foster Sustainable Urban Mobility,\u201d Sustainability, vol. 16, no. 16, p. 6910, 2024, doi: 10.3390/su16166910.",
    "[3] C. Gao and Y. Chen, \u201cUsing Machine Learning Methods to Predict Demand for Bike Sharing,\u201d in Information and Communication Technologies in Tourism 2022 (ENTER 2022), pp. 282-296, 2022, doi: 10.1007/978-3-030-94751-4_25.",
    "[4] M. Schnieder, \u201cEbike Sharing vs. Bike Sharing: Demand Prediction Using Deep Neural Networks and Random Forests,\u201d Sustainability, vol. 15, no. 18, p. 13898, 2023, doi: 10.3390/su151813898.",
    "[5] H. Fanaee-T and J. Gama, \u201cEvent labeling combining ensemble detectors and background knowledge,\u201d Progress in Artificial Intelligence, vol. 2, no. 2-3, pp. 113-127, 2014.",
]
for pk in pustaka:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(pk)
    set_font(r, 12)

doc.add_page_break()

# ============ LAMPIRAN ============
add_heading(doc, "LAMPIRAN")
doc.add_paragraph()
add_subheading(doc, "Lampiran 1. Link Repository GitHub")
p = doc.add_paragraph()
r = p.add_run("https://github.com/alvinalfandy/bike-sharing-analysis")
set_font(r, 12)
r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
doc.add_paragraph()
add_subheading(doc, "Lampiran 2. Link Deploy Aplikasi")
p = doc.add_paragraph()
r = p.add_run("https://bike-sharing-analysis-aby.streamlit.app")
set_font(r, 12)
r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
doc.add_paragraph()
add_subheading(doc, "Lampiran 3. Poster Proyek")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[ Sisipkan Gambar Poster Proyek di sini ]")
set_font(r, 11, italic=True)

doc.save("Laporan_Proyek_Akhir_Data_Mining_2026.docx")
print("Laporan berhasil dibuat: Laporan_Proyek_Akhir_Data_Mining_2026.docx")
